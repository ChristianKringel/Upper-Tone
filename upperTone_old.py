"""
Para rodar este código fora do executável através do próprio python é necessário instalar o docx
pip install python-docx
"""
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import re
from docx import Document

class FileManager:
    file_path = ""
    arquivo_nome = ""

class nmbr_tones:
    nmbr_of_tones = ""

class Screen:
    def __init__(self, master):
        self.Screen = master
        self.botao_clicado = False 

        # Configuração da janela principal
        self.Screen.title("🎵 Upper Tone - Editor de Notas Musicais")
        self.Screen.geometry("600x550")
        self.Screen.configure(bg='#1e1e2e')
        self.Screen.resizable(False, False)

        # Frame principal com padding
        self.frame_principal = tk.Frame(self.Screen, bg='#1e1e2e')
        self.frame_principal.pack(fill='both', expand=True, padx=20, pady=20)

        # Título
        self.titulo = tk.Label(
            self.frame_principal,
            text="🎵 Upper Tone",
            font=('Helvetica', 24, 'bold'),
            bg='#1e1e2e',
            fg='#cdd6f4'
        )
        self.titulo.pack(pady=(0, 5))

        self.subtitulo = tk.Label(
            self.frame_principal,
            text="Transponha suas músicas facilmente",
            font=('Helvetica', 11),
            bg='#1e1e2e',
            fg='#a6adc8'
        )
        self.subtitulo.pack(pady=(0, 30))

        # Seção 1: Selecionar Arquivo
        self.criar_secao_arquivo()

        # Seção 2: Configurar Tons
        self.criar_secao_tons()

        # Seção 3: Processar
        self.criar_secao_processar()

        # Status Bar
        self.status_bar = tk.Label(
            self.Screen,
            text="Pronto para começar",
            font=('Helvetica', 9),
            bg='#313244',
            fg='#cdd6f4',
            anchor='w',
            padx=10
        )
        self.status_bar.pack(side='bottom', fill='x')
        
        # Centraliza a janela após criar todos os widgets
        self.Screen.update_idletasks()
        self.center_window()

    def center_window(self):
        width = 600
        height = 550
        x = (self.Screen.winfo_screenwidth() // 2) - (width // 2)
        y = (self.Screen.winfo_screenheight() // 2) - (height // 2)
        self.Screen.geometry(f'{width}x{height}+{x}+{y}')

    def criar_secao_arquivo(self):
        # Frame da seção
        frame_arquivo = tk.Frame(self.frame_principal, bg='#313244', relief='flat')
        frame_arquivo.pack(fill='x', pady=(0, 20))

        # Padding interno
        frame_interno = tk.Frame(frame_arquivo, bg='#313244')
        frame_interno.pack(fill='both', padx=15, pady=15)

        # Título da seção
        tk.Label(
            frame_interno,
            text="📁 Passo 1: Selecione o Arquivo",
            font=('Helvetica', 13, 'bold'),
            bg='#313244',
            fg='#cdd6f4'
        ).pack(anchor='w', pady=(0, 10))

        # Label de arquivo selecionado
        self.label_arquivo = tk.Label(
            frame_interno,
            text="Nenhum arquivo selecionado",
            font=('Helvetica', 10),
            bg='#313244',
            fg='#a6adc8',
            wraplength=500,
            justify='left'
        )
        self.label_arquivo.pack(anchor='w', pady=(0, 10))

        # Botão para selecionar arquivo
        self.btn_selecionar = tk.Button(
            frame_interno,
            text="📂 Escolher Arquivo",
            command=self.readFile,
            font=('Helvetica', 11, 'bold'),
            bg='#89b4fa',
            fg='#1e1e2e',
            activebackground='#74c7ec',
            activeforeground='#1e1e2e',
            cursor='hand2',
            relief='flat',
            padx=20,
            pady=10
        )
        self.btn_selecionar.pack(anchor='w')
    
    def criar_secao_tons(self):
        # Frame da seção
        frame_tons = tk.Frame(self.frame_principal, bg='#313244', relief='flat')
        frame_tons.pack(fill='x', pady=(0, 20))

        # Padding interno
        frame_interno = tk.Frame(frame_tons, bg='#313244')
        frame_interno.pack(fill='both', padx=15, pady=15)

        # Título da seção
        tk.Label(
            frame_interno,
            text="🎹 Passo 2: Configure a Transposição",
            font=('Helvetica', 13, 'bold'),
            bg='#313244',
            fg='#cdd6f4'
        ).pack(anchor='w', pady=(0, 10))

        # Frame para direção
        frame_direcao = tk.Frame(frame_interno, bg='#313244')
        frame_direcao.pack(fill='x', pady=(0, 10))

        tk.Label(
            frame_direcao,
            text="Direção:",
            font=('Helvetica', 10),
            bg='#313244',
            fg='#cdd6f4'
        ).pack(side='left', padx=(0, 10))

        self.direction_var = tk.StringVar(value="subir")
        
        tk.Radiobutton(
            frame_direcao,
            text="⬆ Subir",
            variable=self.direction_var,
            value="subir",
            font=('Helvetica', 10),
            bg='#313244',
            fg='#cdd6f4',
            selectcolor='#45475a',
            activebackground='#313244',
            activeforeground='#cdd6f4',
            cursor='hand2'
        ).pack(side='left', padx=5)

        tk.Radiobutton(
            frame_direcao,
            text="⬇ Descer",
            variable=self.direction_var,
            value="descer",
            font=('Helvetica', 10),
            bg='#313244',
            fg='#cdd6f4',
            selectcolor='#45475a',
            activebackground='#313244',
            activeforeground='#cdd6f4',
            cursor='hand2'
        ).pack(side='left', padx=5)

        # Frame para quantidade
        frame_qntd = tk.Frame(frame_interno, bg='#313244')
        frame_qntd.pack(fill='x')

        tk.Label(
            frame_qntd,
            text="Quantidade de semitons (1-11):",
            font=('Helvetica', 10),
            bg='#313244',
            fg='#cdd6f4'
        ).pack(anchor='w', pady=(0, 5))

        # Entry estilizado
        self.qntd_entry = tk.Entry(
            frame_qntd,
            font=('Helvetica', 12),
            bg='#45475a',
            fg='#cdd6f4',
            insertbackground='#cdd6f4',
            relief='flat',
            width=10
        )
        self.qntd_entry.pack(anchor='w', ipady=5, ipadx=5)
        self.qntd_entry.insert(0, "1")

    def criar_secao_processar(self):
        # Frame da seção
        frame_processar = tk.Frame(self.frame_principal, bg='#313244', relief='flat')
        frame_processar.pack(fill='x')

        # Padding interno
        frame_interno = tk.Frame(frame_processar, bg='#313244')
        frame_interno.pack(fill='both', padx=15, pady=15)

        # Título da seção
        tk.Label(
            frame_interno,
            text="✨ Passo 3: Processar e Salvar",
            font=('Helvetica', 13, 'bold'),
            bg='#313244',
            fg='#cdd6f4'
        ).pack(anchor='w', pady=(0, 15))

        # Botão processar
        self.btn_processar = tk.Button(
            frame_interno,
            text="🚀 Processar Arquivo",
            command=self.save_and_close,
            font=('Helvetica', 12, 'bold'),
            bg='#a6e3a1',
            fg='#1e1e2e',
            activebackground='#94e2d5',
            activeforeground='#1e1e2e',
            cursor='hand2',
            relief='flat',
            padx=25,
            pady=12,
            state='disabled'
        )
        self.btn_processar.pack(fill='x')
    
    def readFile(self):
        self.arquivo = filedialog.askopenfile(
            mode="rb",
            initialdir=os.path.expanduser("~"),
            title="Selecione um arquivo de música",
            filetypes=(
                ("Arquivos do Word", "*.docx"),
                ("Arquivos de Texto", "*.txt"),
                ("Todos os arquivos", "*.*")
            )
        )
        if self.arquivo:
            FileManager.file_path = self.arquivo.name
            FileManager.arquivo_nome = os.path.basename(self.arquivo.name)
            
            # Atualizar UI
            self.label_arquivo.config(
                text=f"✓ {FileManager.arquivo_nome}",
                fg='#a6e3a1'
            )
            self.btn_processar.config(state='normal')
            self.status_bar.config(text=f"Arquivo carregado: {FileManager.arquivo_nome}")
            self.arquivo.close()

    # Seleciona onde salvar o arquivo e depois encerra a execução
    def save_and_close(self):
        # Validar entrada
        try:
            quantity = int(self.qntd_entry.get())
            if quantity < 1 or quantity > 11:
                messagebox.showerror(
                    "Erro de Validação",
                    "Por favor, insira um número entre 1 e 11."
                )
                return
        except ValueError:
            messagebox.showerror(
                "Erro de Validação",
                "Por favor, insira um número válido."
            )
            return

        # Verificar se arquivo foi selecionado
        if not FileManager.file_path:
            messagebox.showwarning(
                "Arquivo Não Selecionado",
                "Por favor, selecione um arquivo primeiro."
            )
            return

        # Ajustar sinal baseado na direção
        if self.direction_var.get() == "descer":
            quantity = -quantity
        
        nmbr_tones.nmbr_of_tones = str(quantity)

        # Salvar arquivo
        self.arquivo_salvar = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"transposed_{FileManager.arquivo_nome}",
            filetypes=(
                ("Documentos Word", "*.docx"),
                ("Todos os arquivos", "*.*")
            )
        )
        
        if self.arquivo_salvar:
            self.status_bar.config(text="Processando arquivo...")
            self.Screen.update()
            caminho_saida = self.arquivo_salvar
            self.Screen.quit()
            main(caminho_saida)

    # dep 
    def verificarClique(self):
        if self.botao_clicado and nmbr_tones.nmbr_of_tones != "":
            print("O botão foi clicado e a quantidade de tons foi definida.")
            self.Screen.quit()

# faz a subida de nota por nota. 
def up_tone(nota, semitons):
    notas = ["DO", "DO#", "RE", "RE#", "MI", "FA", "FA#", "SOL", "SOL#", "LA", "LA#", "SI"]
    if nota not in notas:
        return nota  # Se a nota não estiver na lista, retorna a nota original (para tratamento de erros)
    index = notas.index(nota)
    new_index = (index + semitons) % len(notas)
    return notas[new_index]

# le o arquivo, sua extensao e caminho 
def read_file(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    linhas = []
    
    try:
        if extension == '.txt':
            with open(file_path, 'r') as file:
                conteudo = file.readlines()  # lê o arquivo linha por linha
        elif extension == '.docx':
            doc = Document(file_path)
            conteudo = [para.text for para in doc.paragraphs]
        else:
            print(f"Extensão de arquivo {extension} não suportada.")
            return []

        for linha in conteudo:
            # regex para separar por espaços ou hífens que não estão no meio das notas
            notas_linha = re.findall(r'[A-Z]+#?', linha.strip().upper())
            delimitadores = re.split(r'[A-Z]+#?', linha.strip().upper())
            linhas.append((notas_linha, delimitadores))
        return linhas
    except FileNotFoundError:
        print(f"Arquivo {file_path} não encontrado.")
        return []

# processa nota por nota, mantendo a formatação do arquivo, como nome, linhas, separação de notas, etc
def process_notes(linhas, semitons):
    new_rows = []
    for notas, delimitadores in linhas:
        notas_modificadas = [up_tone(nota, semitons) for nota in notas]
        nova_linha = ''
        for i in range(len(notas_modificadas)):
            nova_linha += delimitadores[i] + notas_modificadas[i]
        nova_linha += delimitadores[-1]  # Adiciona o último delimitador
        new_rows.append(nova_linha)
    return new_rows


# salva o arquivo no path que o usuário escolhe 
def save_file(file_path, linhas):
    doc = Document()
    for linha in linhas:
        doc.add_paragraph(linha)
    doc.save(file_path)

# printa o resultado apos processamento das notas
def show_result(linhas):
    for linha in linhas:
        print(linha)

def main(caminho_saida):
    file_path = FileManager.file_path
    file_input_path = file_path  # Nome do arquivo de entrada (pode ser .txt ou .docx)
    
    # Verifica se a quantidade de tons foi definida
    if not nmbr_tones.nmbr_of_tones:
        print("Quantidade de tons não definida.")
        return
    
    try:
        semitones = int(nmbr_tones.nmbr_of_tones)
    except ValueError:
        print("Quantidade de tons inválida.")
        return

    # Le as notas do arquivo
    rows = read_file(file_input_path)
    
    if not rows:
        print("Nenhuma nota foi lida do arquivo.")
        return
    
    # Processa notas
    modified_notes = process_notes(rows, semitones)
    
    # Salva e exibe resultado
    save_file(caminho_saida, modified_notes)
    show_result(modified_notes)
    
    # Mostrar mensagem de sucesso
    try:
        root = tk.Tk()
        root.withdraw()
        messagebox.showinfo(
            "Sucesso!",
            f"Arquivo processado com sucesso!\n\nSalvo em:\n{caminho_saida}"
        )
        root.destroy()
    except:
        print(f"Arquivo salvo com sucesso em: {caminho_saida}")

if __name__ == "__main__":
    mainFrame = tk.Tk()
    screen_app = Screen(mainFrame)
    mainFrame.mainloop()
    screen_app.verificarClique()
