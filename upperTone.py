"""
Para rodar este código fora do executável através do próprio python é necessário instalar:
pip install python-docx customtkinter pillow
"""
import customtkinter as ctk
from tkinter import filedialog, messagebox
import os
import re
from docx import Document

# Configurações do CustomTkinter
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class FileManager:
    file_path = ""
    arquivo_nome = ""

class nmbr_tones:
    nmbr_of_tones = ""

class Screen:
    def __init__(self, master):
        self.Screen = master
        self.botao_clicado = False 

        # Configuração da janela principal
        self.Screen.title("Upper Tone - Editor de Notas Musicais")
        self.Screen.geometry("700x720")
        self.Screen.resizable(False, False)

        # Frame principal com padding
        self.frame_principal = ctk.CTkFrame(
            self.Screen,
            fg_color="transparent"
        )
        self.frame_principal.pack(fill='both', expand=True, padx=25, pady=25)

        # Header com título
        self.criar_header()

        # Seção 1: Selecionar Arquivo
        self.criar_secao_arquivo()

        # Seção 2: Configurar Tons
        self.criar_secao_tons()

        # Seção 3: Processar
        self.criar_secao_processar()

        # Status Bar
        self.criar_status_bar()

    def criar_header(self):
        # Frame do header
        header_frame = ctk.CTkFrame(
            self.frame_principal,
            fg_color="transparent"
        )
        header_frame.pack(fill='x', pady=(0, 25))

        # Título principal
        titulo = ctk.CTkLabel(
            header_frame,
            text="Upper Tone",
            font=ctk.CTkFont(family="Helvetica", size=36, weight="bold"),
            text_color=("#1f6aa5", "#3b8ed0")
        )
        titulo.pack(pady=(0, 5))

        # Subtítulo
        subtitulo = ctk.CTkLabel(
            header_frame,
            text="Transponha suas músicas de forma rápida e profissional",
            font=ctk.CTkFont(size=14),
            text_color="gray70"
        )
        subtitulo.pack()

    def criar_secao_arquivo(self):
        # Card da seção
        card_frame = ctk.CTkFrame(
            self.frame_principal,
            corner_radius=15,
            fg_color=("#e0e0e0", "#2b2b2b")
        )
        card_frame.pack(fill='x', pady=(0, 15))

        # Conteúdo interno
        content = ctk.CTkFrame(card_frame, fg_color="transparent")
        content.pack(fill='both', padx=20, pady=20)

        # Título da seção com badge
        header_section = ctk.CTkFrame(content, fg_color="transparent")
        header_section.pack(fill='x', pady=(0, 15))

        badge = ctk.CTkLabel(
            header_section,
            text="1",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color="white",
            fg_color=("#1f6aa5", "#1f6aa5"),
            corner_radius=12,
            width=28,
            height=28
        )
        badge.pack(side='left', padx=(0, 10))

        titulo_secao = ctk.CTkLabel(
            header_section,
            text="Selecione o Arquivo",
            font=ctk.CTkFont(size=18, weight="bold"),
            anchor='w'
        )
        titulo_secao.pack(side='left', fill='x', expand=True)

        # Label de arquivo selecionado
        self.label_arquivo = ctk.CTkLabel(
            content,
            text="Nenhum arquivo selecionado",
            font=ctk.CTkFont(size=13),
            text_color="gray60",
            anchor='w'
        )
        self.label_arquivo.pack(fill='x', pady=(0, 12))

        # Botão para selecionar arquivo
        self.btn_selecionar = ctk.CTkButton(
            content,
            text="Escolher Arquivo Musical",
            command=self.readFile,
            font=ctk.CTkFont(size=14, weight="bold"),
            height=42,
            corner_radius=10,
            fg_color=("#1f6aa5", "#1f6aa5"),
            hover_color=("#144870", "#144870"),
            cursor='hand2'
        )
        self.btn_selecionar.pack(fill='x')
    
    def criar_secao_tons(self):
        # Card da seção
        card_frame = ctk.CTkFrame(
            self.frame_principal,
            corner_radius=15,
            fg_color=("#e0e0e0", "#2b2b2b")
        )
        card_frame.pack(fill='x', pady=(0, 15))

        # Conteúdo interno
        content = ctk.CTkFrame(card_frame, fg_color="transparent")
        content.pack(fill='both', padx=20, pady=20)

        # Título da seção com badge
        header_section = ctk.CTkFrame(content, fg_color="transparent")
        header_section.pack(fill='x', pady=(0, 15))

        badge = ctk.CTkLabel(
            header_section,
            text="2",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color="white",
            fg_color=("#1f6aa5", "#1f6aa5"),
            corner_radius=12,
            width=28,
            height=28
        )
        badge.pack(side='left', padx=(0, 10))

        titulo_secao = ctk.CTkLabel(
            header_section,
            text="Configure a Transposição",
            font=ctk.CTkFont(size=18, weight="bold"),
            anchor='w'
        )
        titulo_secao.pack(side='left', fill='x', expand=True)

        # Frame para direção
        direcao_label = ctk.CTkLabel(
            content,
            text="Direção da Transposição:",
            font=ctk.CTkFont(size=13, weight="bold"),
            anchor='w'
        )
        direcao_label.pack(fill='x', pady=(0, 8))

        self.direction_var = ctk.StringVar(value="subir")
        
        radio_frame = ctk.CTkFrame(content, fg_color="transparent")
        radio_frame.pack(fill='x', pady=(0, 15))

        radio_subir = ctk.CTkRadioButton(
            radio_frame,
            text="Subir Tom",
            variable=self.direction_var,
            value="subir",
            font=ctk.CTkFont(size=13),
            radiobutton_width=20,
            radiobutton_height=20,
            cursor='hand2'
        )
        radio_subir.pack(side='left', padx=(0, 20))

        radio_descer = ctk.CTkRadioButton(
            radio_frame,
            text="Descer Tom",
            variable=self.direction_var,
            value="descer",
            font=ctk.CTkFont(size=13),
            radiobutton_width=20,
            radiobutton_height=20,
            cursor='hand2'
        )
        radio_descer.pack(side='left')

        # Frame para quantidade
        qntd_label = ctk.CTkLabel(
            content,
            text="Quantidade de Semitons:",
            font=ctk.CTkFont(size=13, weight="bold"),
            anchor='w'
        )
        qntd_label.pack(fill='x', pady=(0, 8))

        # Slider moderno
        slider_frame = ctk.CTkFrame(content, fg_color="transparent")
        slider_frame.pack(fill='x')

        self.slider_value = ctk.CTkLabel(
            slider_frame,
            text="1",
            font=ctk.CTkFont(size=28, weight="bold"),
            text_color=("#1f6aa5", "#3b8ed0"),
            width=60
        )
        self.slider_value.pack(side='right', padx=(15, 0))

        self.slider = ctk.CTkSlider(
            slider_frame,
            from_=1,
            to=11,
            number_of_steps=10,
            command=self.update_slider_value,
            width=380,
            height=20,
            button_color=("#1f6aa5", "#1f6aa5"),
            button_hover_color=("#144870", "#144870"),
            progress_color=("#1f6aa5", "#1f6aa5")
        )
        self.slider.set(1)
        self.slider.pack(side='left', fill='x', expand=True)

    def update_slider_value(self, value):
        self.slider_value.configure(text=str(int(value)))

    def criar_secao_processar(self):
        # Card da seção
        card_frame = ctk.CTkFrame(
            self.frame_principal,
            corner_radius=15,
            fg_color=("#e0e0e0", "#2b2b2b")
        )
        card_frame.pack(fill='x', pady=(0, 15))

        # Conteúdo interno
        content = ctk.CTkFrame(card_frame, fg_color="transparent")
        content.pack(fill='both', padx=20, pady=20)

        # Título da seção com badge
        header_section = ctk.CTkFrame(content, fg_color="transparent")
        header_section.pack(fill='x', pady=(0, 15))

        badge = ctk.CTkLabel(
            header_section,
            text="3",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color="white",
            fg_color=("#1f6aa5", "#1f6aa5"),
            corner_radius=12,
            width=28,
            height=28
        )
        badge.pack(side='left', padx=(0, 10))

        titulo_secao = ctk.CTkLabel(
            header_section,
            text="Processar e Salvar",
            font=ctk.CTkFont(size=18, weight="bold"),
            anchor='w'
        )
        titulo_secao.pack(side='left', fill='x', expand=True)

        # Botão processar grande e destacado
        self.btn_processar = ctk.CTkButton(
            content,
            text="🚀 Processar Arquivo",
            command=self.save_and_close,
            font=ctk.CTkFont(size=16, weight="bold"),
            height=50,
            corner_radius=12,
            fg_color=("#16a34a", "#16a34a"),
            hover_color=("#15803d", "#15803d"),
            cursor='hand2',
            state='disabled'
        )
        self.btn_processar.pack(fill='x', pady=(0, 0))

    def criar_status_bar(self):
        # Status bar moderno
        self.status_frame = ctk.CTkFrame(
            self.frame_principal,
            corner_radius=8,
            height=35,
            fg_color=("#c0c0c0", "#1a1a1a")
        )
        self.status_frame.pack(fill='x')
        
        self.status_bar = ctk.CTkLabel(
            self.status_frame,
            text="Pronto para começar",
            font=ctk.CTkFont(size=12),
            text_color="gray70",
            anchor='w'
        )
        self.status_bar.pack(fill='both', padx=15, pady=6)
    
    def readFile(self):
        arquivo = filedialog.askopenfile(
            mode="rb",
            initialdir=os.path.expanduser("~"),
            title="Selecione um arquivo de música",
            filetypes=(
                ("Arquivos do Word", "*.docx"),
                ("Arquivos de Texto", "*.txt"),
                ("Todos os arquivos", "*.*")
            )
        )
        if arquivo:
            FileManager.file_path = arquivo.name
            FileManager.arquivo_nome = os.path.basename(arquivo.name)
            
            # Atualizar UI com animação
            self.label_arquivo.configure(
                text=f"✓ Arquivo selecionado: {FileManager.arquivo_nome}",
                text_color=("#16a34a", "#22c55e")
            )
            self.btn_processar.configure(
                state='normal',
                text="🚀 Processar Arquivo (Clique Aqui!)"
            )
            self.status_bar.configure(
                text=f"✓ Arquivo carregado: {FileManager.arquivo_nome}",
                text_color=("#16a34a", "#22c55e")
            )
            arquivo.close()
            print(f"[DEBUG] Arquivo carregado: {FileManager.arquivo_nome}")
            print(f"[DEBUG] Botão processar habilitado")

    def save_and_close(self):
        print("[DEBUG] save_and_close chamado!")
        
        # Validar entrada
        try:
            quantity = int(self.slider.get())
            print(f"[DEBUG] Quantidade selecionada: {quantity}")
            if quantity < 1 or quantity > 11:
                messagebox.showerror(
                    "Erro de Validação",
                    "Por favor, selecione um valor entre 1 e 11."
                )
                return
        except ValueError:
            messagebox.showerror(
                "Erro de Validação",
                "Por favor, selecione um número válido."
            )
            return

        # Verificar se arquivo foi selecionado
        if not FileManager.file_path:
            messagebox.showwarning(
                "Arquivo Não Selecionado",
                "Por favor, selecione um arquivo primeiro."
            )
            return

        print(f"[DEBUG] Arquivo de entrada: {FileManager.file_path}")
        
        # Ajustar sinal baseado na direção
        direcao = self.direction_var.get()
        print(f"[DEBUG] Direção: {direcao}")
        
        if direcao == "descer":
            quantity = -quantity
        
        nmbr_tones.nmbr_of_tones = str(quantity)
        print(f"[DEBUG] Semitons finais: {quantity}")

        # Salvar arquivo
        arquivo_salvar = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"transposed_{FileManager.arquivo_nome}",
            filetypes=(
                ("Documentos Word", "*.docx"),
                ("Todos os arquivos", "*.*")
            )
        )
        
        if arquivo_salvar:
            print(f"[DEBUG] Arquivo de saída: {arquivo_salvar}")
            self.status_bar.configure(
                text="⏳ Processando arquivo...",
                text_color=("#d97706", "#f59e0b")
            )
            self.btn_processar.configure(
                state='disabled',
                text="⏳ Processando..."
            )
            self.Screen.update()
            caminho_saida = arquivo_salvar
            self.Screen.quit()
            main(caminho_saida)
        else:
            print("[DEBUG] Usuário cancelou o salvamento")

    def verificarClique(self):
        if self.botao_clicado and nmbr_tones.nmbr_of_tones != "":
            print("O botão foi clicado e a quantidade de tons foi definida.")
            self.Screen.quit()

# faz a subida de nota por nota
def up_tone(nota, semitons):
    notas = ["DO", "DO#", "RE", "RE#", "MI", "FA", "FA#", "SOL", "SOL#", "LA", "LA#", "SI"]
    if nota not in notas:
        return nota
    index = notas.index(nota)
    new_index = (index + semitons) % len(notas)
    return notas[new_index]

# le o arquivo, sua extensao e caminho
def read_file(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    linhas = []
    
    try:
        if extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                conteudo = file.readlines()
        elif extension == '.docx':
            doc = Document(file_path)
            conteudo = [para.text for para in doc.paragraphs]
        else:
            print(f"Extensão de arquivo {extension} não suportada.")
            return []

        for linha in conteudo:
            notas_linha = re.findall(r'[A-Z]+#?', linha.strip().upper())
            delimitadores = re.split(r'[A-Z]+#?', linha.strip().upper())
            linhas.append((notas_linha, delimitadores))
        return linhas
    except FileNotFoundError:
        print(f"Arquivo {file_path} não encontrado.")
        return []

# processa nota por nota
def process_notes(linhas, semitons):
    new_rows = []
    for notas, delimitadores in linhas:
        notas_modificadas = [up_tone(nota, semitons) for nota in notas]
        nova_linha = ''
        for i in range(len(notas_modificadas)):
            nova_linha += delimitadores[i] + notas_modificadas[i]
        nova_linha += delimitadores[-1]
        new_rows.append(nova_linha)
    return new_rows

# salva o arquivo
def save_file(file_path, linhas):
    doc = Document()
    for linha in linhas:
        doc.add_paragraph(linha)
    doc.save(file_path)

# mostra resultado
def show_result(linhas):
    for linha in linhas:
        print(linha)

def main(caminho_saida):
    file_path = FileManager.file_path
    file_input_path = file_path
    
    if not nmbr_tones.nmbr_of_tones:
        print("Quantidade de tons não definida.")
        return
    
    try:
        semitones = int(nmbr_tones.nmbr_of_tones)
    except ValueError:
        print("Quantidade de tons inválida.")
        return

    rows = read_file(file_input_path)
    
    if not rows:
        print("Nenhuma nota foi lida do arquivo.")
        return
    
    modified_notes = process_notes(rows, semitones)
    save_file(caminho_saida, modified_notes)
    show_result(modified_notes)
    
    # Mensagem de sucesso com CustomTkinter
    try:
        dialog = ctk.CTkToplevel()
        dialog.title("Sucesso!")
        dialog.geometry("500x200")
        dialog.resizable(False, False)
        
        # Centralizar
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - 250
        y = (dialog.winfo_screenheight() // 2) - 100
        dialog.geometry(f"500x200+{x}+{y}")
        
        frame = ctk.CTkFrame(dialog, fg_color="transparent")
        frame.pack(fill='both', expand=True, padx=30, pady=30)
        
        titulo = ctk.CTkLabel(
            frame,
            text="Arquivo Processado com Sucesso!",
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=("#16a34a", "#22c55e")
        )
        titulo.pack(pady=(0, 15))
        
        mensagem = ctk.CTkLabel(
            frame,
            text=f"Salvo em:\n{caminho_saida}",
            font=ctk.CTkFont(size=12),
            wraplength=440
        )
        mensagem.pack(pady=(0, 20))
        
        btn_ok = ctk.CTkButton(
            frame,
            text="OK",
            command=dialog.destroy,
            width=120,
            height=35,
            font=ctk.CTkFont(size=14, weight="bold")
        )
        btn_ok.pack()
        
        dialog.transient()
        dialog.grab_set()
        dialog.wait_window()
    except:
        print(f"Arquivo salvo com sucesso em: {caminho_saida}")

if __name__ == "__main__":
    app = ctk.CTk()
    screen_app = Screen(app)
    app.mainloop()
    screen_app.verificarClique()
